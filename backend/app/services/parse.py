from __future__ import annotations
import re
from pathlib import Path
from typing import Tuple
import fitz  # pymupdf
from docx import Document

SECTION_HINTS = [
    "summary", "skills", "experience", "education", "projects", "certifications",
    "achievements", "publications", "internship"
]

def _clean_text(t: str) -> str:
    t = t.replace("\x00", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()

def extract_text_from_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    chunks = []
    for page in doc:
        chunks.append(page.get_text("text"))
    return _clean_text("\n".join(chunks))

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # include tables (basic)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return _clean_text("\n".join(parts))

def detect_formatting_flags(ext: str, raw_text: str) -> dict:
    t = raw_text.lower()
    has_email = bool(re.search(r"[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}", t))
    has_phone = bool(re.search(r"(\+?\d[\d \-\(\)]{8,}\d)", t))
    has_linkedin = "linkedin.com" in t

    # Heuristic: very short lines might suggest columns (common in PDFs)
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    short_line_ratio = 0.0
    if lines:
        short_line_ratio = sum(1 for ln in lines if len(ln) <= 25) / max(len(lines), 1)

    section_hits = [s for s in SECTION_HINTS if s in t]

    flags = {
        "file_type": ext,
        "contact_info": {
            "email_detected": has_email,
            "phone_detected": has_phone,
            "linkedin_detected": has_linkedin,
        },
        "possible_multi_column_layout": (ext == ".pdf" and short_line_ratio > 0.45),
        "section_presence": {
            "detected_sections": section_hits,
            "missing_core_sections": [s for s in ["summary", "skills", "experience", "education"] if s not in section_hits],
        },
        "readability": {
            "line_count": len(lines),
            "short_line_ratio": round(short_line_ratio, 3),
        },
    }
    return flags

def parse_resume(file_path: str) -> Tuple[str, dict]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type")

    flags = detect_formatting_flags(ext, text)
    return text, flags
